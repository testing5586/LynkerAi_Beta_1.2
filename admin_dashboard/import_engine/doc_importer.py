"""
DOC/DOCX 文档格式导入器
Word Document Importer - 使用 python-docx 解析 Word 文档
"""

import os
from supabase import create_client
from .normalize_chart import normalize_from_ocr

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
sb = create_client(SUPABASE_URL, SUPABASE_KEY) if SUPABASE_URL and SUPABASE_KEY else None

# 动态加载 python-docx（可选依赖）
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def extract_text_from_docx(file_bytes):
    """从 DOCX 文件中提取全部文本"""
    if not DOCX_AVAILABLE:
        return None, "python-docx 未安装，请运行: pip install python-docx"
    
    try:
        from io import BytesIO
        doc = Document(BytesIO(file_bytes))
        
        # 提取所有段落文本
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text), None
        
    except Exception as e:
        return None, f"DOCX 解析失败: {str(e)}"

def process_docx_file(file_bytes):
    """
    处理 DOCX 文件
    Returns: (parsed_data, error_message)
    """
    if not DOCX_AVAILABLE:
        return {
            "error": "DOC/DOCX 功能未安装",
            "message": "请运行: pip install python-docx",
            "name": None,
            "gender": None,
            "birth_time": None,
            "marriage_palace_star": None,
            "wealth_palace_star": None,
            "transformations": {"hualu": False, "huaji": False},
            "raw_text": "DOC/DOCX 功能未启用"
        }, "python-docx 未安装"
    
    # 提取文本
    text, error = extract_text_from_docx(file_bytes)
    
    if error:
        return {
            "error": error,
            "name": None,
            "gender": None,
            "birth_time": None,
            "marriage_palace_star": None,
            "wealth_palace_star": None,
            "transformations": {"hualu": False, "huaji": False},
            "raw_text": ""
        }, error
    
    # 使用 TXT 解析器处理提取的文本
    from .txt_importer import extract_fields_from_text
    parsed = extract_fields_from_text(text)
    
    return parsed, None

def save_record(doc_obj):
    """保存 DOC/DOCX 解析后的记录到数据库"""
    if not sb:
        raise Exception("Supabase 客户端未初始化")
    
    norm = normalize_from_ocr(doc_obj)
    
    return sb.table("birthcharts").insert({
        "name": norm["name"],
        "gender": norm["gender"],
        "birth_time": norm["birth_time"],
        "ziwei_palace": None,
        "main_star": None,
        "shen_palace": None,
        "birth_data": norm["birth_data"]
    }).execute()
